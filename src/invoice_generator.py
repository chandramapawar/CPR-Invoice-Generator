from dataclasses import dataclass
from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from num2words import num2words

@dataclass
class Record:
    course_name: str
    officer_name: str
    rank: str
    duration_days: int
    from_date: date | None = None
    to_date: date | None = None
    police_unit: str = ""
    reference_no: str = ""
    fee_per_day: int = 2000

    @property
    def total_fee(self):
        return self.duration_days * self.fee_per_day

class InvoiceGenerator:
    def fmt_money(self, x):
        return f"{int(x):,}" if float(x).is_integer() else f"{x:,.2f}"

    def _set_font(self, run, name="Calibri", size=11, bold=False, underline=False):
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        run.font.size = Pt(size)
        run.bold = bold
        run.underline = underline

    def _para(self, parent, text="", size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = parent.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        self._set_font(r, size=size, bold=bold)
        return p

    def _set_cell(self, cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        self._set_font(r, size=size, bold=bold)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _set_cell_margins(self, cell, top=40, start=40, bottom=40, end=40):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = tcPr.first_child_found_in("w:tcMar")
        if tcMar is None:
            tcMar = OxmlElement("w:tcMar")
            tcPr.append(tcMar)
        for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
            node = tcMar.find(qn(f"w:{m}"))
            if node is None:
                node = OxmlElement(f"w:{m}")
                tcMar.append(node)
            node.set(qn("w:w"), str(v))
            node.set(qn("w:type"), "dxa")

    def amount_in_words(self, n):
        return num2words(n).replace("-", " ").title()

    def generate_invoice(self, records, output_path, invoice_no, invoice_date):
        first = records[0]
        total = sum(r.total_fee for r in records)
        unit_name = first.police_unit.strip()
        ref = f"OW/CPR/{unit_name}/        /2026"
        inv_str = f"CPR/INV/2026/{invoice_no}"

        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Cm(1.0)
        sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2)
        sec.right_margin = Cm(1.2)

        self._para(doc, "Centre for Police Research, Pune", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._para(doc, "Chavan Nagar, Pashan Road, Pune, Maharashtra - 411008", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._para(doc, "Email: directorcprpune@gmail.com, Phone: 020-25653696", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._para(doc, "GST No. -- 27AAATC3424CIZB PAN - AAATC3424C", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._para(doc, "Demand Note/Invoice", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        hdr = doc.add_table(rows=1, cols=2)
        hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr.autofit = True
        self._set_cell(
            hdr.cell(0, 0),
            f"Invoice  No: {inv_str}",
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        self._set_cell(
            hdr.cell(0, 1),
            f"Date: {invoice_date.strftime('%d-%m-%Y')}",
            size=11,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        for c in hdr.rows[0].cells:
            self._set_cell_margins(c)

        self._para(doc, f"Reference: {ref}", size=10)
        self._para(doc, "To,", size=10)
        self._para(doc, "Superintendent of Police,", size=10)
        self._para(doc, first.police_unit, size=10)
        self._para(doc, "Training Details", size=11, bold=True)

        table = doc.add_table(rows=1, cols=9)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Sr No", "Training Name", "From Date", "To Date", "Officer Name", "Rank", "Duration (Days)", "Fee/Day", "Total Fee"]
        for i, h in enumerate(headers):
            self._set_cell(table.rows[0].cells[i], h, size=9, bold=True)

        for i, r in enumerate(records, 1):
            row = table.add_row().cells
            self._set_cell(row[0], str(i), size=9)
            self._set_cell(row[1], r.course_name, size=9)
            self._set_cell(row[2], r.from_date.strftime("%d/%m/%y") if r.from_date else "", size=9)
            self._set_cell(row[3], r.to_date.strftime("%d/%m/%y") if r.to_date else "", size=9)
            self._set_cell(row[4], r.officer_name, size=9)
            self._set_cell(row[5], r.rank, size=9)
            self._set_cell(row[6], str(r.duration_days), size=9)
            self._set_cell(row[7], "₹2000", size=9)
            self._set_cell(row[8], f"₹{self.fmt_money(r.total_fee)}", size=9)

        gp = doc.add_paragraph()
        gp.paragraph_format.space_before = Pt(4)
        gp.paragraph_format.space_after = Pt(0)
        rr = gp.add_run(f"Grand Total: ₹{self.fmt_money(total)}")
        self._set_font(rr, bold=True)

        self._para(doc, f"Amount: {self.amount_in_words(int(total))} Rupees Only", size=10)

        self._para(doc, "Bank Account Details", size=11, bold=True)
        bank_items = [
            "Account Name: CPR, Pune",
            "Account Number: 10023971530",
            "Type: Savings",
            "Bank Name: State Bank of India",
            "Branch Address: NCL Branch, Pashan Road, Pune - 411008",
            "Branch Code: 3552",
            "MICR Code: 411002012",
            "IFSC Code: SBIN0003552",
            "Payee Code: 22010023828",
        ]
        for item in bank_items:
            self._para(doc, item, size=10)

        sig = doc.add_paragraph()
        sig.paragraph_format.space_before = Pt(4)
        sig.paragraph_format.space_after = Pt(0)
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = sig.add_run("(Dr. Kakasaheb Dole)\nSuperintendent of Police\nCentre for Police Research, Pune")
        self._set_font(rr, size=11)

        doc.save(output_path)
        return output_path