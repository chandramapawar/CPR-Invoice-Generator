from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class LetterGenerator:
    def fmt_money(self, x):
        return f"{int(x):,}" if float(x).is_integer() else f"{x:,.2f}"

    def fmt_money_marathi(self, x):
        eng = self.fmt_money(x)
        mapping = str.maketrans("0123456789", "०१२३४५६७८९")
        return eng.translate(mapping)

    def _marathi_unit(self, name):
        m = {
            # Core districts / cities
            "Amravati City": "अमरावती शहर",
            "Amravati": "अमरावती",
            "Ahilyanagar":"अहिल्यानगर",
            "Aurangabad": "छत्रपती संभाजीनगर",
            "Chhatrapati Sambhajinagar": "छत्रपतीसंभाजीनगर",
            "Chhatrapati Sambhajinagar City": "छत्रपती संभाजीनगर शहर",
            "Dhule": "धुळे",
            "Gadchiroli": "गडचिरोली",
            "Hingoli": "हिंगोली",
            "Jalgaon": "जळगाव",
            "Kolhapur": "कोल्हापूर",
            "Latur": "लातूर",
            "Nagpur": "नागपूर",
            "Nagpur City": "नागपूर शहर",
            "Nanded": "नांदेड",
            "Nashik": "नाशिक",
            "Navi Mumbai": "नवीमुंबई",
            "Palghar": "पालघर",
            "Parbhani": "परभणी",
            "Pune": "पुणे",
            "Pune City": "पुणे शहर",
            "Raigad": "रायगड",
            "Ratnagiri": "रत्नागिरी",
            "Sangli": "सांगली",
            "Satara": "सातारा",
            "Sindhudurg": "सिंधुदुर्ग",
            "Solapur": "सोलापूर",
            "Thane": "ठाणे",
            "Wardha": "वर्धा",
            "Washim": "वाशीम",
            "Yavatmal": "यवतमाळ",

            # Commissionerates / metros
            "Amaravati": "अमरावती",
            "Brihan Mumbai": "बृहन्मुंबई",
            "Chhatrapati Sambhaji Nagar": "छत्रपतीसंभाजीनगर",
            "Mira-Bhayandar, Vasai-Virar": "मीरा-भाईंदर,वसई-विरार",
            "Pimpri Chinchwad": "पिंपरीचिंचवड",
            "Railways Mumbai": "रेल्वेमुंबई",

            # Districts / Rural variants
            "Ahmednagar": "अहमदनगर",
            "Akola": "अकोला",
            "Amaravati R": "अमरावतीग्रामीण",
            "Beed": "बीड",
            "Bhandara": "भंडारा",
            "Buldhana": "बुलढाणा",
            "Chandrapur": "चंद्रपूर",
            "Chhatrapati Sambhaji Nagar Rural": "छत्रपतीसंभाजीनगरग्रामीण",
            "Dharashiv": "धाराशिव",
            "Gondia": "गोंदिया",
            "Jalna": "जालना",
            "Nagpur R": "नागपूरग्रामीण",
            "Nandurbar": "नंदुरबार",
            "Nashik R": "नाशिकग्रामीण",
            "Pune R": "पुणेग्रामीण",
            "Solapur R": "सोलापूरग्रामीण",
            "Thane R": "ठाणेग्रामीण",
            "Railways Nagpur": "रेल्वेनागपूर",
            "Railways Pune": "रेल्वेपुणे",
            "Railways Chhatrapati Sambhaji Nagar": "रेल्वेछत्रपतीसंभाजीनगर",

            # PTCs
            "PTC Dhule": "पोलीसप्रशिक्षणकेंद्रधुळे",
            "PTC Jalna": "पोलीसप्रशिक्षणकेंद्रजालना",
            "PTC Khandala": "पोलीसप्रशिक्षणकेंद्रखंडाळा",
            "PTC Marol": "पोलीसप्रशिक्षणकेंद्रमरोळ",
            "PTC Nanvij": "पोलीसप्रशिक्षणकेंद्रनानवीज",
            "PTC Solapur": "पोलीसप्रशिक्षणकेंद्रसोलापूर",
            "PTC Turchi": "पोलीसप्रशिक्षणकेंद्रतुर्ची",
            "PTC Akola": "पोलीसप्रशिक्षणकेंद्रअकोला",
            "PTC Latur": "पोलीसप्रशिक्षणकेंद्रलातूर",
            "PTC Nagpur": "पोलीसप्रशिक्षणकेंद्रनागपूर",

            # SRPF Groups
            "SRPF Gr 1 Pune": "राज्यराखीवपोलीसबलगटक्रमांक१पुणे",
            "SRPF Gr 2 Pune": "राज्यराखीवपोलीसबलगटक्रमांक२पुणे",
            "SRPF Gr 3 Jalna": "राज्यराखीवपोलीसबलगटक्रमांक३जालना",
            "SRPF Gr 4 Nagpur": "राज्यराखीवपोलीसबलगटक्रमांक४नागपूर",
            "SRPF Gr 5 Daund": "राज्यराखीवपोलीसबलगटक्रमांक५दौंड",
            "SRPF Gr 6 Dhule": "राज्यराखीवपोलीसबलगटक्रमांक६धुळे",
            "SRPF Gr 7 Daund": "राज्यराखीवपोलीसबलगटक्रमांक७दौंड",
            "SRPF Gr 8 Mumbai": "राज्यराखीवपोलीसबलगटक्रमांक८मुंबई",
            "SRPF Gr 9 Amravati": "राज्यराखीवपोलीसबलगटक्रमांक९अमरावती",
            "SRPF Gr 10 Solapur": "राज्यराखीवपोलीसबलगटक्रमांक१०सोलापूर",
            "SRPF Gr 11 Mumbai": "राज्यराखीवपोलीसबलगटक्रमांक११मुंबई",
            "SRPF Gr 12 Hingoli": "राज्यराखीवपोलीसबलगटक्रमांक१२हिंगोली",
            "SRPF Gr 13 Gadchiroli": "राज्य राखीव पोलीस बल गट क्र. १३ गडचिरोली",
            "SRPF Gr 14 Sambhajinagar": "राज्य राखीव पोलीस बल गट क्र. १४ संभाजीनगर",
            "SRPF Gr 15 Gondia": "राज्य राखीव पोलीस बल गट क्र. १५ गोंदिया",
            "SRPF Gr 16 Kolhapur": "राज्य राखीव पोलीस बल गट क्र. १६ कोल्हापूर",

            # Special units
            "Mahamarg Raigad": "महामार्ग रायगड",
            "SID": "राज्य गुप्तवार्ता विभाग",
            "ACB PUNE": "लाचलुचपत प्रतिबंधक विभाग",
            "ACB Thane": "लाचलुचपत प्रतिबंधक विभाग",
            "PCR": "नागरी हक्क संरक्षण",
            "ATS": "दहशतवाद विरोधी पथक",
            "ATS Pune": "दहशतवाद विरोधी पथक",
            "Nashik MPA": "महाराष्ट्र पोलीस प्रबोधिनी",
            "Force One": "फोर्स वन",
            "CID": "गुन्हे अन्वेषण विभाग",
        }
        n = (name or "").strip()
        return m.get(n, n)

    def _set_font(self, run, name="Mangal", size=12, bold=False, underline=False):
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        run.font.size = Pt(size)
        run.bold = bold
        run.underline = underline

    def _set_normal_style(self, doc):
        style = doc.styles["Normal"]
        style.font.name = "Mangal"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Mangal")
        style.font.size = Pt(12)

    def _center_para(self, doc, text, size=12, bold=False, underline=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        self._set_font(r, size=size, bold=bold, underline=underline)
        return p

    def _set_cell(self, cell, text, name="Mangal", size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        self._set_font(r, name=name, size=size, bold=bold)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _remove_table_borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.first_child_found_in("w:tblBorders")
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = tblBorders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tblBorders.append(element)
            element.set(qn("w:val"), "nil")

    def generate_letter(self, records, output_path, reference_no, letter_date):
        first = records[0]
        unit_name = self._marathi_unit(getattr(first, "police_unit", "") or "")

        date_str = letter_date.strftime("%d/%m/%Y")
        marathi_date = date_str.translate(str.maketrans("0123456789", "०१२३४५६७८९"))

        doc = Document()
        self._set_normal_style(doc)

        sec = doc.sections[0]
        sec.top_margin = Cm(1.0)
        sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.2)
        sec.right_margin = Cm(1.2)

        self._center_para(doc, "पोलीस संशोधन केंद्र, पुणे", size=14, bold=True)
        self._center_para(doc, "चव्हाण नगर, पाषाण रोड, पुणे ४११००८", size=11)
        self._center_para(doc, "TEL NO- 020-25653696 FAX NO- 020-25653696", size=10)
        self._center_para(doc, "E-MAIL- directorcprpune@gmail.com WEBSITE- www.cprpune.org", size=10)

        hdr = doc.add_table(rows=1, cols=2)
        hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr.autofit = False
        hdr.columns[0].width = Cm(11.5)
        hdr.columns[1].width = Cm(5.5)
        self._remove_table_borders(hdr)
        self._set_cell(
            hdr.cell(0, 0),
            "जा.क्र.सी.पी.आर./प्रलंबित प्रशिक्षण शुल्क/         /२०२६",
            size=11,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )
        self._set_cell(
            hdr.cell(0, 1),
            f"पुणे दि. {marathi_date}",
            size=11,
            bold=False,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run("प्रति,")

        doc.add_paragraph("मा. पोलीस अधिक्षक,")

        unit_p = doc.add_paragraph()
        unit_p.paragraph_format.space_before = Pt(0)
        unit_p.paragraph_format.space_after = Pt(0)
        unit_p.add_run(unit_name)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run("विषय")
        self._set_font(r1, bold=True, underline=True)
        r2 = p.add_run(" --- पोलीस संशोधन केंद्र, पुणे येथील आर्थिक वर्ष २०२६-२०२७ मधील प्रलंबित प्रशिक्षण शुल्क अदा करणेबाबत.")
        self._set_font(r2)

        doc.add_paragraph("महोदय,")

        body = doc.add_paragraph()
        body.paragraph_format.space_after = Pt(0)
        body.add_run(
            "उपरोक्त विषय व संदर्भान्वये अनुसरून पोलीस संशोधन केंद्र, पुणे येथे मा. पोलीस महासंचालक, महाराष्ट्र राज्य, मुंबई यांच्या निर्देशानुसार महाराष्ट्र पोलीस दलातील पोलीस अधीक्षक ते पोलीस उप निरीक्षक दर्जाचे अधिकारी यांचे करीता नियमित सीपीआर येथे वेगवेगळ्या विषयांवर पोलीस सेवेमधील अधिका-यांसाठी प्रशिक्षण आयोजित होत असतात. आपल्या घटकातील अधिका-यांनी आर्थिक वर्ष २०२६-२०२७ मधील खालील तक्त्यातील विषयांवर प्रशिक्षण घेतलेले असून त्यांचे प्रशिक्षण रक्कम येणे बाकी आहे."
        )

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["अ.क्र", "प्रशिक्षणाचे नाव व सत्र", "प्रशिक्षणाचा कालावधी", "प्रशिक्षणार्थींचे नाव", "प्रशिक्षण रक्कम", "इनव्हॉइस नंबर"]
        for i, h in enumerate(headers):
            self._set_cell(table.rows[0].cells[i], h, size=10, bold=True)

        for i, r in enumerate(records, 1):
            row = table.add_row().cells
            self._set_cell(row[0], str(i), size=10)
            self._set_cell(row[1], f"{r.course_name} {r.batch_session}".strip(), size=10)
            self._set_cell(
                row[2],
                f"{r.from_date.strftime('%d/%m/%Y') if r.from_date else ''} ते {r.to_date.strftime('%d/%m/%Y') if r.to_date else ''}",
                size=10,
            )
            self._set_cell(row[3], r.officer_name, size=10)
            self._set_cell(row[4], self.fmt_money(r.total_fee), size=10)
            self._set_cell(row[5], str(reference_no).strip(), size=10)

        doc.add_paragraph(
            "सदर प्रशिक्षणासाठी आपले घटकाकडून वरील प्रमाणे थकबाकी रक्कम अद्याप पोलीस संशोधन केंद्र पुणे या कार्यालयास प्राप्त झालेली नाही."
        )

        total = sum(r.total_fee for r in records)

        doc.add_paragraph(
            f"प्रती प्रशिक्षणार्थी प्रती दिवस रू.२०००/- या प्रमाणे एकूण रक्कम रू. {self.fmt_money_marathi(total)}/- या प्रमाणे प्रशिक्षण शुल्क पोलीस संशोधन केंद्र पुणे (CPR Pune) या नावाने धनादेश अथवा खालील बँक खात्यामध्ये RTGS / NEFT ने जमा करण्यास आणि त्याची माहिती directorcprpune@gmail.com या ई-मेलवर पाठविण्याची विनंती आहे."
        )

        doc.add_paragraph(
            "तसेच वर नमुद पोलीस अधिकारी यांचे प्रशिक्षण शुल्क यापूर्वी NEFT/RTGS/ धनादेश ई. व्दारे जमा केले असल्यास UTR क्रमांक, दिनांक, पोलीस अधिकारी यांचे नाव प्रशिक्षण शुल्क ई. माहिती directorcprpune@gmail.com या ई-मेलवर तात्काळ पाठविण्याची विनंती आहे."
        )

        doc.add_paragraph(
            "तरी सदर थकीत प्रशिक्षण फी रक्कम पोलीस संशोधन केंद्र पुणे येथे तातडीने पाठविण्याची विनंती आहे."
        )

        doc.add_paragraph(
            "सोबतः- वर नमुद इन व्हॉईस नंबर प्रमाणे प्रशिक्षण फी Demand Note/Invoice जोडलेले आहे."
        )

        bank = doc.add_table(rows=1, cols=3)
        bank.style = "Table Grid"
        bank.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_cell(bank.rows[0].cells[0], "अ.क्र", size=10, bold=True)
        self._set_cell(bank.rows[0].cells[1], "Field", size=10, bold=True)
        self._set_cell(bank.rows[0].cells[2], "Value", size=10, bold=True)

        bank_rows = [
            ("1", "Name", "CPR, Pune"),
            ("2", "Account No", "10023971530"),
            ("3", "Type of Account", "Saving Bank Account"),
            ("4", "Bank Name", "State Bank of India"),
            ("5", "Bank Address", "NCL Branch, Pashan Road, Pune 411008."),
            ("6", "Bank Branch Code", "3552"),
            ("7", "MICR Code", "411002012"),
            ("8", "IFSC Code", "SBIN0003552"),
            ("9", "Mobile No.", "9960631393"),
            ("10", "Payee Code", "22010023828"),
        ]
        for no, field, val in bank_rows:
            row = bank.add_row().cells
            self._set_cell(row[0], no, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell(row[1], field, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
            self._set_cell(row[2], val, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

        sig = doc.add_paragraph()
        sig.paragraph_format.space_before = Pt(4)
        sig.paragraph_format.space_after = Pt(0)
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig.add_run("(डॉ. काकासाहेब डोळे)\nपोलीस अधीक्षक,\nपोलीस संशोधन केंद्र पुणे")

        doc.save(output_path)
        return output_path